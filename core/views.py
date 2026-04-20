import asyncio
import edge_tts
import os
from faster_whisper import WhisperModel

# small модель — жылдам (mic, audio-text үшін)
_whisper_model = None
# large-v3 модель — дәл (субтитр үшін)
_whisper_large_model = None


def get_whisper_model():
    global _whisper_model
    if _whisper_model is None:
        cpu_count = os.cpu_count() or 4
        _whisper_model = WhisperModel(
            "small",
            device="cpu",
            compute_type="int8",
            cpu_threads=cpu_count,
            num_workers=2,
        )
    return _whisper_model


def get_whisper_large_model():
    """medium модель — қазақша субтитрге арналған (large-v3-тен 3x жылдам, дәлдік жақын)"""
    global _whisper_large_model
    if _whisper_large_model is None:
        cpu_count = os.cpu_count() or 4
        for model_name in ("medium", "large-v3"):
            try:
                _whisper_large_model = WhisperModel(
                    model_name,
                    device="cpu",
                    compute_type="int8",
                    cpu_threads=cpu_count,
                    num_workers=2,
                )
                print(f"✅ Субтитр моделі жүктелді: {model_name}")
                break
            except Exception as e:
                print(f"⚠️ {model_name} жүктелмеді: {e}, келесіге өту...")
    return _whisper_large_model


def _transcribe_fast(audio_path, language=None):
    """Жылдам транскрипция — mic/audio_text үшін"""
    model = get_whisper_model()
    segments_gen, info = model.transcribe(
        audio_path,
        language=language,
        beam_size=1,
        best_of=1,
        temperature=0.0,
        condition_on_previous_text=False,
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500),
    )
    return segments_gen, info


def _transcribe_accurate(audio_path, language=None):
    """Дәл транскрипция — субтитр үшін large-v3 моделі"""
    model = get_whisper_large_model()

    # Қазақша үшін контекст үлгісі
    prompts = {
        "kk": "Қазақ тіліндегі дәріс. Сөздер дұрыс жазылуы керек.",
        "ru": "Лекция на русском языке.",
        "en": "This is a lecture in English.",
    }
    initial_prompt = prompts.get(language)

    segments_gen, info = model.transcribe(
        audio_path,
        language=language,
        beam_size=3,
        best_of=3,
        temperature=[0.0, 0.2],
        condition_on_previous_text=True,
        initial_prompt=initial_prompt,
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=400),
        word_timestamps=False,
        no_speech_threshold=0.6,
        log_prob_threshold=-1.0,
        compression_ratio_threshold=2.4,
    )
    return segments_gen, info
os.environ["PATH"] = r"E:\ffmpeg-8.1-essentials_build\bin" + os.pathsep + os.environ.get("PATH", "")

os.environ['TESSDATA_PREFIX'] = r'E:\tess\tessdata'
import fitz  # PyMuPDF
import whisper
import pytesseract
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm
from reportlab.lib import colors

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login
from django.http import FileResponse, Http404, HttpResponse
from .models import ConvertedFile
from django.contrib.auth.models import User
from django.contrib import messages
from django.db.models import Count
import django, sys, os
from datetime import datetime, timedelta
from django.shortcuts import render, redirect, get_object_or_404



# Tesseract жолын көрсет (Windows үшін)
pytesseract.pytesseract.tesseract_cmd = r'E:\tess\tesseract.exe'



# ─────────────────────────────────────────────
# 1. TTS — Мәтін → Аудио
# ─────────────────────────────────────────────
async def generate_voice(text, output_path):
    """Мәтінді MP3-ке айналдыру — edge_tts 3000 таңбадан асса бөліп жіберу"""
    MAX_CHARS = 3000

    if len(text) <= MAX_CHARS:
        communicate = edge_tts.Communicate(text, "kk-KZ-AigulNeural")
        await communicate.save(output_path)
        return

    # Үлкен мәтінді сөйлемдерге бөліп, параллель генерация
    import re, tempfile
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, current = [], ""
    for s in sentences:
        if len(current) + len(s) < MAX_CHARS:
            current += s + " "
        else:
            if current:
                chunks.append(current.strip())
            current = s + " "
    if current:
        chunks.append(current.strip())

    # Барлық чанкты параллель генерация
    tmp_files = []
    async def gen_chunk(chunk, idx):
        tmp = output_path + f".part{idx}.mp3"
        communicate = edge_tts.Communicate(chunk, "kk-KZ-AigulNeural")
        await communicate.save(tmp)
        return tmp

    tasks = [gen_chunk(c, i) for i, c in enumerate(chunks)]
    tmp_files = await asyncio.gather(*tasks)

    # ffmpeg арқылы біріктіру
    list_file = output_path + ".list.txt"
    with open(list_file, "w", encoding="utf-8") as f:
        for t in tmp_files:
            f.write(f"file '{os.path.abspath(t)}'\n")

    os.system(f'ffmpeg -y -f concat -safe 0 -i "{list_file}" -c copy "{output_path}" -loglevel error')

    # Уақытша файлдарды тазалау
    for t in tmp_files:
        try: os.remove(t)
        except: pass
    try: os.remove(list_file)
    except: pass


def convert_tts(myfile, user):
    """PDF/DOCX → MP3 аудио"""
    os.makedirs('media/uploads', exist_ok=True)
    os.makedirs('media/audio', exist_ok=True)

    doc = fitz.open(stream=myfile.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()

    if not full_text.strip():
        return None, "PDF-тен мәтін табылмады"

    audio_name = f"tts_{os.path.splitext(myfile.name)[0]}.mp3"
    audio_path = os.path.join('media/audio', audio_name)
    asyncio.run(generate_voice(full_text, audio_path))

    new_file = ConvertedFile.objects.create(
        user=user,
        title=myfile.name,
        original_file=myfile,
        converted_audio=f"audio/{audio_name}",
        conversion_type='tts'
    )
    return new_file.converted_audio.url, None


# ─────────────────────────────────────────────
# 2. SUBTITLES — Видео → Субтитрлер
# ─────────────────────────────────────────────

def format_time(seconds):
    """Секундты SRT уақыт форматына айналдыру"""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

def format_time_vtt(seconds):
    """Секундты VTT уақыт форматына айналдыру"""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"

def convert_subtitles(myfile, user, language=None):
    """MP4 видео → SRT + VTT субтитр. language=None болса авто-анықтайды."""
    try:
        os.makedirs('media/uploads', exist_ok=True)
        os.makedirs('media/subtitles', exist_ok=True)

        # Видеоны сақтау
        video_path = os.path.join('media/uploads', myfile.name)
        with open(video_path, 'wb+') as f:
            for chunk in myfile.chunks():
                f.write(chunk)

        print(f"✅ Видео сақталды: {video_path}")
        print("⏳ Whisper модель жүктелуде...")
        get_whisper_model()
        print("✅ Модель дайын")

        # Дәл транскрипция — субтитр үшін жоғары сапа
        forced_lang = language if language and language != 'auto' else None
        print(f"🔍 Транскрипция басталды (тіл: {forced_lang or 'авто'})...")
        segments_gen, info = _transcribe_accurate(video_path, language=forced_lang)
        detected_lang = info.language or 'unknown'
        segments = list(segments_gen)
        print(f"✅ Анықталған тіл: {detected_lang}")
        print(f"✅ Транскрипция: {len(segments)} сегмент")
        # segments → result dict форматына дайындау
        result = {
            'language': detected_lang,
            'segments': [{'start': s.start, 'end': s.end, 'text': s.text} for s in segments],
        }

        base_name = os.path.splitext(myfile.name)[0]

        # SRT жасау
        srt_name = f"{base_name}.srt"
        srt_path = os.path.join('media/subtitles', srt_name)
        with open(srt_path, 'w', encoding='utf-8') as f:
            for i, segment in enumerate(result['segments'], 1):
                start = format_time(segment['start'])
                end = format_time(segment['end'])
                text = segment['text'].strip()
                f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

        # VTT жасау (браузер үшін)
        vtt_name = f"{base_name}.vtt"
        vtt_path = os.path.join('media/subtitles', vtt_name)
        with open(vtt_path, 'w', encoding='utf-8') as f:
            f.write("WEBVTT\n\n")
            for i, segment in enumerate(result['segments'], 1):
                start = format_time_vtt(segment['start'])
                end = format_time_vtt(segment['end'])
                text = segment['text'].strip()
                f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

        # Тіл атауы
        lang_names = {
            'kk': 'Қазақша', 'ru': 'Орысша', 'en': 'Ағылшынша',
            'unknown': 'Белгісіз'
        }
        lang_name = lang_names.get(detected_lang, detected_lang.upper())

        # Субтитрді видеоға ендіру (ffmpeg)
        import subprocess
        os.makedirs('media/subtitles', exist_ok=True)
        burned_name = f"{base_name}_subtitled.mp4"
        burned_path = os.path.join('media/subtitles', burned_name)

        # Windows-та subtitles фильтрі үшін жолды дұрыс экрандау:
        # E:\foo\bar.srt → E\\:/foo/bar.srt
        srt_abs = os.path.abspath(srt_path).replace('\\', '/')
        if len(srt_abs) > 1 and srt_abs[1] == ':':
            srt_abs = srt_abs[0] + '\\\\:' + srt_abs[2:]
        vf_filter = (
            f"subtitles='{srt_abs}'"
            ":force_style='FontSize=20,PrimaryColour=&H00FFFFFF,"
            "OutlineColour=&H00000000,Outline=2,Bold=1'"
        )
        ffmpeg_proc = subprocess.run(
            [
                'ffmpeg', '-y',
                '-i', os.path.abspath(video_path),
                '-vf', vf_filter,
                '-c:a', 'copy',
                '-loglevel', 'error',
                os.path.abspath(burned_path),
            ],
            capture_output=True, text=True
        )
        if ffmpeg_proc.returncode != 0:
            print(f"ffmpeg қатесі: {ffmpeg_proc.stderr}")
        burned_url = (
            f"/media/subtitles/{burned_name}"
            if ffmpeg_proc.returncode == 0 and os.path.exists(burned_path)
            else None
        )

        ConvertedFile.objects.create(
            user=user,
            title=myfile.name,
            conversion_type='subtitles',
            converted_subtitle=f"subtitles/{srt_name}"
        )
        return {
            'video_url': f"/media/uploads/{myfile.name}",
            'srt_url': f"/media/subtitles/{srt_name}",
            'vtt_url': f"/media/subtitles/{vtt_name}",
            'burned_url': burned_url,
            'burned_name': burned_name,
            'srt_name': srt_name,
            'detected_lang': lang_name,
        }, None

    except Exception as e:
        print(f"❌ ҚАТЕ: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return None, str(e)





# ─────────────────────────────────────────────
# 3. OCR — Сурет → Мәтін
# ─────────────────────────────────────────────
def convert_ocr(myfile, user):
    """Сурет (PNG/JPG) немесе PDF → Мәтін"""
    os.makedirs('media/uploads', exist_ok=True)
    os.makedirs('media/ocr', exist_ok=True)

    file_ext = os.path.splitext(myfile.name)[1].lower()
    extracted_text = ""

    if file_ext == '.pdf':
        # PDF-тен сурет ретінде оқу
        doc = fitz.open(stream=myfile.read(), filetype="pdf")
        for page in doc:
            pix = page.get_pixmap()
            img_path = f"media/uploads/temp_page.png"
            pix.save(img_path)
            img = Image.open(img_path)
            extracted_text += pytesseract.image_to_string(img, lang='rus+kaz+eng') + "\n"
    else:
        # Сурет файлы
        img = Image.open(myfile)
        extracted_text = pytesseract.image_to_string(img, lang='rus+kaz+eng')

    if not extracted_text.strip():
        return None, "Суреттен мәтін табылмады"

    # Мәтінді файлға сақтау
    txt_name = f"ocr_{os.path.splitext(myfile.name)[0]}.txt"
    txt_path = os.path.join('media/ocr', txt_name)
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(extracted_text)

    new_file = ConvertedFile.objects.create(
        user=user,
        title=myfile.name,
        conversion_type='ocr',
        converted_text=extracted_text[:500]  # Алғашқы 500 символды сақтау
    )
    return extracted_text, f"/media/ocr/{txt_name}", None


# ─────────────────────────────────────────────
# 4. LARGE TEXT — Үлкен оқылатын мәтін
# ─────────────────────────────────────────────
def convert_large_text(myfile, user, font_size=20, color_scheme='white',
                        line_spacing=1.5, font_type='Arial', page_size='A4'):
    """PDF/DOCX/TXT/сурет → тақырып иерархиясы, кестелер, суреттер сақталған үлкен шрифтті PDF"""
    import io, statistics
    from reportlab.lib.pagesizes import A4, A3
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Image as RLImage, Table as RLTable,
                                    TableStyle, PageBreak, HRFlowable)
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    os.makedirs('media/large_text', exist_ok=True)

    ext = os.path.splitext(myfile.name)[1].lower()

    # ── Шрифт ────────────────────────────────────────────────────────────────
    font_name = 'Helvetica'
    font_bold  = 'Helvetica-Bold'
    font_candidates = {
        'Arial':        [(r'C:\Windows\Fonts\arial.ttf',   r'C:\Windows\Fonts\arialbd.ttf')],
        'OpenDyslexic': [(r'C:\Windows\Fonts\OpenDyslexic-Regular.otf', None)],
    }
    for reg_path, bold_path in font_candidates.get(font_type, font_candidates['Arial']):
        if os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont('CustomFont', reg_path))
                font_name = 'CustomFont'
                if bold_path and os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('CustomFont-Bold', bold_path))
                    font_bold = 'CustomFont-Bold'
                else:
                    font_bold = font_name
                break
            except Exception:
                pass
    if font_name == 'Helvetica':
        for reg_path, bold_path in [
            (r'C:\Windows\Fonts\arial.ttf',   r'C:\Windows\Fonts\arialbd.ttf'),
            (r'C:\Windows\Fonts\calibri.ttf', r'C:\Windows\Fonts\calibrib.ttf'),
        ]:
            if os.path.exists(reg_path):
                try:
                    pdfmetrics.registerFont(TTFont('CustomFont', reg_path))
                    font_name = 'CustomFont'
                    if bold_path and os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont('CustomFont-Bold', bold_path))
                        font_bold = 'CustomFont-Bold'
                    else:
                        font_bold = font_name
                    break
                except Exception:
                    pass

    # ── Түс схемасы ──────────────────────────────────────────────────────────
    color_map = {
        'white':  (colors.white,               colors.black,               colors.HexColor('#1a56db')),
        'black':  (colors.black,               colors.white,               colors.HexColor('#60a5fa')),
        'yellow': (colors.HexColor('#FFFDE7'), colors.black,               colors.HexColor('#1a56db')),
        'cream':  (colors.HexColor('#FFF8E1'), colors.HexColor('#222222'), colors.HexColor('#1a3a6b')),
    }
    bg_color, text_color, heading_color = color_map.get(
        color_scheme, color_map['white'])

    # ── Бет өлшемі ───────────────────────────────────────────────────────────
    selected_page_size = A3 if page_size == 'A3' else A4
    page_w, _ = selected_page_size
    usable_w  = page_w - 4 * cm

    # ── Стильдер ─────────────────────────────────────────────────────────────
    fs  = int(font_size)
    ls  = fs * float(line_spacing)

    def _style(name, size, bold=False, color=None, space_before=0, space_after=4):
        return ParagraphStyle(name,
            fontSize=size, leading=size * float(line_spacing),
            textColor=color or text_color,
            fontName=font_bold if bold else font_name,
            spaceBefore=space_before, spaceAfter=space_after,
            wordWrap='CJK', alignment=TA_LEFT)

    body_style = _style('Body', fs, space_after=fs // 2)
    h1_style   = _style('H1',  int(fs * 1.6), bold=True,  color=heading_color,
                         space_before=fs, space_after=fs // 2)
    h2_style   = _style('H2',  int(fs * 1.35), bold=True, color=heading_color,
                         space_before=int(fs * 0.8), space_after=fs // 3)
    h3_style   = _style('H3',  int(fs * 1.15), bold=True, color=heading_color,
                         space_before=int(fs * 0.6), space_after=fs // 4)
    cell_style = _style('Cell', max(fs - 2, 8), space_after=2)

    # ── Шығыс PDF ────────────────────────────────────────────────────────────
    pdf_name = f"large_{os.path.splitext(myfile.name)[0]}_{fs}pt_{color_scheme}.pdf"
    pdf_path  = os.path.join('media/large_text', pdf_name)

    def make_bg(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFillColor(bg_color)
        canvas_obj.rect(0, 0, *doc_obj.pagesize, fill=1, stroke=0)
        canvas_obj.restoreState()

    def safe_para(txt, style):
        txt = (txt or '').strip().replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        try:
            return Paragraph(txt, style) if txt else None
        except Exception:
            return None

    def make_rl_table(raw_rows):
        """raw_rows → RLTable; None қайтарады егер жасалмаса"""
        if not raw_rows:
            return None
        num_cols = max(len(r) for r in raw_rows)
        if num_cols == 0:
            return None
        col_w = usable_w / num_cols
        rl_data = []
        for row in raw_rows:
            rl_row = [safe_para(str(cell or ''), cell_style) or
                      Paragraph('', cell_style) for cell in row]
            # Жол ұяшықтарын тегестіру
            while len(rl_row) < num_cols:
                rl_row.append(Paragraph('', cell_style))
            rl_data.append(rl_row)
        tbl = RLTable(rl_data, colWidths=[col_w] * num_cols, splitByRow=True)
        tbl.setStyle(TableStyle([
            ('GRID',          (0,0), (-1,-1), 0.5, text_color),
            ('BACKGROUND',    (0,0), (-1,-1), bg_color),
            ('TEXTCOLOR',     (0,0), (-1,-1), text_color),
            ('FONTNAME',      (0,0), (-1,-1), font_name),
            ('FONTSIZE',      (0,0), (-1,-1), max(fs-2, 8)),
            ('VALIGN',        (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING',    (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING',   (0,0), (-1,-1), 4),
            ('RIGHTPADDING',  (0,0), (-1,-1), 4),
        ]))
        tbl._raw_fallback = raw_rows
        return tbl

    def add_table(story, raw_rows):
        tbl = make_rl_table(raw_rows)
        if tbl:
            story.append(tbl)
            story.append(Spacer(1, 0.3 * cm))

    def add_image(story, img_bytes):
        try:
            pil_img = Image.open(io.BytesIO(img_bytes))
            orig_w, orig_h = pil_img.size
            if orig_w == 0:
                return
            draw_w = min(usable_w, orig_w * 0.75)
            draw_h = draw_w * (orig_h / orig_w)
            # Бет биіктігінен аспасын
            max_h = selected_page_size[1] - 5 * cm
            if draw_h > max_h:
                draw_h = max_h
                draw_w = draw_h * (orig_w / orig_h)
            story.append(RLImage(io.BytesIO(img_bytes), width=draw_w, height=draw_h))
            story.append(Spacer(1, 0.3 * cm))
        except Exception:
            pass

    # ════════════════════════════════════════════════════════════════════════
    # ФОРМАТҚА БАЙЛАНЫСТЫ МАЗМҰН АЛУ
    # ════════════════════════════════════════════════════════════════════════
    story = []
    file_bytes = myfile.read()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if ext == '.pdf':
        doc_in = fitz.open(stream=file_bytes, filetype='pdf')

        # Денелік шрифт өлшемін анықтау (медиана)
        all_sizes = []
        for pg in doc_in:
            for blk in pg.get_text('dict')['blocks']:
                if blk['type'] == 0:
                    for ln in blk['lines']:
                        for sp in ln['spans']:
                            if sp['text'].strip():
                                all_sizes.append(round(sp['size'], 1))
        body_fs = statistics.median(all_sizes) if all_sizes else 10

        for page_num, page in enumerate(doc_in):
            if page_num > 0:
                story.append(PageBreak())

            table_rects, page_tables = [], []
            try:
                for tbl in page.find_tables():
                    page_tables.append(tbl)
                    table_rects.append(fitz.Rect(tbl.bbox))
            except Exception:
                pass

            def in_table(bbox):
                return any(fitz.Rect(bbox).intersects(tr) for tr in table_rects)

            elements = [(tbl.bbox[1], 'table', tbl) for tbl in page_tables]

            for blk in page.get_text('dict')['blocks']:
                bbox = blk['bbox']
                if in_table(bbox):
                    continue
                if blk['type'] == 1:
                    img_b = blk.get('image')
                    if img_b:
                        elements.append((bbox[1], 'image', img_b))
                elif blk['type'] == 0:
                    # Тақырып деңгейін шрифт өлшемі арқылы анықтау
                    blk_sizes = [sp['size'] for ln in blk['lines']
                                 for sp in ln['spans'] if sp['text'].strip()]
                    blk_fs = max(blk_sizes) if blk_sizes else body_fs
                    ratio  = blk_fs / body_fs if body_fs else 1

                    lines_txt = []
                    for ln in blk['lines']:
                        lt = ' '.join(sp['text'] for sp in ln['spans']).strip()
                        if lt:
                            lines_txt.append(lt)
                    if lines_txt:
                        elements.append((bbox[1], 'block',
                                         ('\n'.join(lines_txt), ratio)))

            elements.sort(key=lambda e: e[0])

            for _, etype, edata in elements:
                if etype == 'block':
                    txt, ratio = edata
                    if ratio >= 1.6:
                        sty = h1_style
                    elif ratio >= 1.25:
                        sty = h2_style
                    elif ratio >= 1.1:
                        sty = h3_style
                    else:
                        sty = body_style
                    for chunk in txt.split('\n\n'):
                        p = safe_para(chunk, sty)
                        if p:
                            story.append(p)
                elif etype == 'image':
                    add_image(story, edata)
                elif etype == 'table':
                    raw = edata.extract()
                    if raw:
                        add_table(story, raw)

    # ── DOCX ─────────────────────────────────────────────────────────────────
    elif ext in ('.docx',):
        from docx import Document as DocxDoc
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxPara

        docx_doc = DocxDoc(io.BytesIO(file_bytes))

        def iter_body(doc):
            for child in doc.element.body:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'p':
                    yield ('para', DocxPara(child, doc))
                elif tag == 'tbl':
                    yield ('table', DocxTable(child, doc))

        for item_type, item in iter_body(docx_doc):
            if item_type == 'para':
                txt  = item.text.strip()
                name = item.style.name if item.style else ''
                if not txt:
                    story.append(Spacer(1, 0.2 * cm))
                    continue
                if 'Heading 1' in name or name.startswith('Title'):
                    sty = h1_style
                elif 'Heading 2' in name:
                    sty = h2_style
                elif 'Heading' in name:
                    sty = h3_style
                else:
                    sty = body_style
                p = safe_para(txt, sty)
                if p:
                    story.append(p)

            elif item_type == 'table':
                raw = []
                for row in item.rows:
                    raw.append(['\n'.join(p.text for p in cell.paragraphs
                                          if p.text.strip())
                                 for cell in row.cells])
                add_table(story, raw)

        # Суреттерді DOCX-тан алу
        for rel in docx_doc.part.rels.values():
            if 'image' in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    add_image(story, img_bytes)
                except Exception:
                    pass

    # ── TXT ──────────────────────────────────────────────────────────────────
    elif ext == '.txt':
        text = file_bytes.decode('utf-8', errors='replace')
        for para in text.split('\n\n'):
            para = para.strip()
            if not para:
                continue
            # Жоғарғы регистрдегі қысқа жол → тақырып
            if para.isupper() and len(para) < 120:
                sty = h2_style
            elif para.startswith('#'):
                sty = h1_style
                para = para.lstrip('#').strip()
            else:
                sty = body_style
            p = safe_para(para, sty)
            if p:
                story.append(p)
                story.append(Spacer(1, 0.2 * cm))

    # ── Сурет (PNG/JPG/BMP/TIFF) ─────────────────────────────────────────────
    elif ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif'):
        add_image(story, file_bytes)

    else:
        return None, f"Қолдаусыз формат: {ext}"

    if not story:
        return None, "Файлдан мазмұн табылмады"

    # ── PDF жасау + fallback ──────────────────────────────────────────────────
    def build_story(s):
        d = SimpleDocTemplate(pdf_path, pagesize=selected_page_size,
                              leftMargin=2*cm, rightMargin=2*cm,
                              topMargin=2*cm, bottomMargin=2*cm)
        d.build(s, onFirstPage=make_bg, onLaterPages=make_bg)

    try:
        build_story(story)
    except Exception:
        fallback = []
        for item in story:
            if isinstance(item, RLTable):
                raw = getattr(item, '_raw_fallback', None)
                if raw:
                    for row in raw:
                        p = safe_para('  |  '.join(str(c or '') for c in row), cell_style)
                        if p:
                            fallback.append(p)
                    fallback.append(Spacer(1, 0.3 * cm))
            else:
                fallback.append(item)
        build_story(fallback)

    ConvertedFile.objects.create(
        user=user, title=myfile.name,
        conversion_type='large_text',
        converted_large_pdf=f"large_text/{pdf_name}"
    )
    return f"/media/large_text/{pdf_name}", None


# ─────────────────────────────────────────────
# 5. BRAILLE — Мәтін → Брайль таңбалары
# ─────────────────────────────────────────────

# Брайль Unicode кестесі (6 нүктелі жүйе)
_BRAILLE_MAP = {
    # Латын әріптері (Grade 1 Unified English Braille)
    'a': '\u2801', 'b': '\u2803', 'c': '\u2809', 'd': '\u2819', 'e': '\u2811',
    'f': '\u280b', 'g': '\u281b', 'h': '\u2813', 'i': '\u280a', 'j': '\u281a',
    'k': '\u2805', 'l': '\u2807', 'm': '\u280d', 'n': '\u281d', 'o': '\u2815',
    'p': '\u280f', 'q': '\u281f', 'r': '\u2817', 's': '\u280e', 't': '\u281e',
    'u': '\u2825', 'v': '\u2827', 'w': '\u283a', 'x': '\u282d', 'y': '\u283d',
    'z': '\u2835',
    # Кириллица (орыс Брайль стандарты ГОСТ 7.86-2003)
    'а': '\u2801', 'б': '\u2803', 'в': '\u283a', 'г': '\u281b', 'д': '\u2819',
    'е': '\u2811', 'ё': '\u2821', 'ж': '\u281a', 'з': '\u2835', 'и': '\u280a',
    'й': '\u281f', 'к': '\u2805', 'л': '\u2807', 'м': '\u280d', 'н': '\u281d',
    'о': '\u2815', 'п': '\u280f', 'р': '\u2817', 'с': '\u280e', 'т': '\u281e',
    'у': '\u2825', 'ф': '\u280b', 'х': '\u2813', 'ц': '\u2809', 'ч': '\u2827',
    'ш': '\u2831', 'щ': '\u2829', 'ъ': '\u282f', 'ы': '\u283d', 'ь': '\u283e',
    'э': '\u282a', 'ю': '\u2833', 'я': '\u282b',
    # Қазақша арнайы таңбалар
    'ə': '\u2823', 'ә': '\u2823',
    'ғ': '\u2838', 'қ': '\u282c', 'ң': '\u2839',
    'ө': '\u2818', 'ұ': '\u2830', 'ү': '\u2837',
    'і': '\u2814', 'ҥ': '\u2839',
    # Сандар (сан белгісі ⠼ алдына қойылады)
    '0': '\u283c\u281a', '1': '\u283c\u2801', '2': '\u283c\u2803',
    '3': '\u283c\u2809', '4': '\u283c\u2819', '5': '\u283c\u2811',
    '6': '\u283c\u280b', '7': '\u283c\u281b', '8': '\u283c\u2813',
    '9': '\u283c\u280a',
    # Тыныс белгілері
    ' ': '\u2800', '\t': '\u2800\u2800', '\n': '\n',
    '.': '\u2832', ',': '\u2802', '!': '\u2816', '?': '\u2826',
    ':': '\u2812', ';': '\u2806', '-': '\u2824', '–': '\u2824\u2824',
    '(': '\u2823', ')': '\u281c', '"': '\u2836', "'": '\u2804',
    '/': '\u280c', '«': '\u2836', '»': '\u2836',
}


def text_to_braille(text):
    """Мәтінді Брайль Unicode таңбаларына айналдыру"""
    result = []
    for char in text:
        lower = char.lower()
        if lower in _BRAILLE_MAP:
            result.append(_BRAILLE_MAP[lower])
        else:
            result.append('\u2800')  # Белгісіз таңба → бос ұяшық
    return ''.join(result)


def convert_braille(text_input, myfile, user):
    """Мәтін немесе PDF → Брайль Unicode (.txt файлы)"""
    os.makedirs('media/braille', exist_ok=True)

    # Мәтінді анықтау: тікелей енгізу немесе PDF-тен
    source_name = 'text_input'
    if text_input:
        raw_text = text_input
        source_name = 'input'
    elif myfile:
        ext = os.path.splitext(myfile.name)[1].lower()
        source_name = os.path.splitext(myfile.name)[0]
        if ext == '.pdf':
            doc = fitz.open(stream=myfile.read(), filetype="pdf")
            raw_text = ''
            for page in doc:
                raw_text += page.get_text()
        else:
            return None, None, "Тек PDF немесе тікелей мәтін енгізіледі"
        if not raw_text.strip():
            return None, None, "PDF-тен мәтін табылмады"
    else:
        return None, None, "Мәтін немесе PDF файл енгізіңіз"

    braille_text = text_to_braille(raw_text)

    txt_name = f"braille_{source_name}.txt"
    txt_path = os.path.join('media/braille', txt_name)
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(braille_text)

    ConvertedFile.objects.create(
        user=user,
        title=source_name,
        conversion_type='braille',
        converted_text=braille_text[:1000],
    )
    return braille_text, f"/media/braille/{txt_name}", None


# ─────────────────────────────────────────────
# VIEWS
# ─────────────────────────────────────────────
def home(request):
    result = {}

    if request.method == 'POST':
        if not request.user.is_authenticated:
            return redirect('login')

        myfile = request.FILES.get('myfile')
        conversion_type = request.POST.get('conversion_type', 'tts')

        try:
            if conversion_type == 'braille':
                braille_text_input = request.POST.get('braille_text', '').strip()
                braille_text, file_url, error = convert_braille(
                    braille_text_input, myfile, request.user
                )
                if braille_text:
                    result = {
                        'type': 'braille',
                        'braille_text': braille_text,
                        'file_url': file_url,
                    }
                else:
                    result = {'error': error}

            elif not myfile:
                result = {'error': 'Файл таңдаңыз'}

            elif conversion_type == 'tts':
                url, error = convert_tts(myfile, request.user)
                if url:
                    result = {'type': 'tts', 'audio_url': url}
                else:
                    result = {'error': error}

            elif conversion_type == 'subtitles':
                subtitle_lang = request.POST.get('subtitle_lang', 'auto')
                data, error = convert_subtitles(myfile, request.user, language=subtitle_lang)
                if data:
                   result = {'type': 'subtitles', **data}
                else:
                     result = {'error': error}

            elif conversion_type == 'ocr':
                text, file_url, error = convert_ocr(myfile, request.user)
                if text:
                    result = {'type': 'ocr', 'ocr_text': text,
                              'file_url': file_url}
                else:
                    result = {'error': error}

            elif conversion_type == 'audio_text':
                text, file_url, lang, error = convert_audio_to_text(myfile, request.user)
                if text:
                    result = {'type': 'audio_text', 'text': text,
                              'file_url': file_url, 'language': lang}
                else:
                    result = {'error': error}

            elif conversion_type == 'large_text':
                font_size = request.POST.get('font_size', '20')
                color_scheme = request.POST.get('color_scheme', 'white')
                line_spacing = request.POST.get('line_spacing', '1.5')
                font_type = request.POST.get('font_type', 'Arial')
                page_size = request.POST.get('page_size', 'A4')
  
                url, error = convert_large_text(
                   myfile, request.user,
                   font_size=font_size,
                   color_scheme=color_scheme,
                   line_spacing=line_spacing,
                   font_type=font_type,
                   page_size=page_size
                )
                if url:
                   result = {'type': 'large_text', 'file_url': url,
                  'filename': os.path.basename(url)}
                else:
                   result = {'error': error}
 
        except Exception as e:
            result = {'error': f"Қате орын алды: {str(e)}"}

    return render(request, 'core/index.html', {'result': result})


@login_required
def delete_user_file(request, file_id):
    if request.method == 'POST':
        file_obj = get_object_or_404(ConvertedFile, id=file_id, user=request.user)
        file_obj.delete()
        messages.success(request, "Файл жойылды.")
    return redirect('profile')


@login_required
def download_text_file(request, file_id):
    """ConvertedFile.converted_text → жүктеп алу"""
    import json
    from django.http import HttpResponse
    file_obj = get_object_or_404(ConvertedFile, id=file_id, user=request.user)
    text = file_obj.converted_text or ''
    ext_map = {'ocr': 'txt', 'braille': 'txt', 'audio_text': 'txt'}
    ext = ext_map.get(file_obj.conversion_type, 'txt')
    filename = f"{file_obj.conversion_type}_{file_obj.title}.{ext}"
    response = HttpResponse(text, content_type='text/plain; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def braille_download(request, filename):
    """Брайль файлын жүктеп алу"""
    file_path = os.path.join('media/braille', filename)
    if not os.path.exists(file_path):
        raise Http404("Файл табылмады")
    response = HttpResponse(
        open(file_path, 'rb').read(),
        content_type='text/plain; charset=utf-8'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def signup(request):
    from .models import UserProfile
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            disability_type = request.POST.get('disability_type', 'visual')
            if disability_type not in ('visual', 'hearing'):
                disability_type = 'visual'
            UserProfile.objects.create(user=user, disability_type=disability_type)
            login(request, user)
            return redirect('home')
    else:
        form = UserCreationForm()
    return render(request, 'registration/signup.html', {'form': form})


@login_required
def profile(request):
    from .models import UserMessage

    user_files = ConvertedFile.objects.filter(user=request.user).order_by('-created_at')
    user_messages = UserMessage.objects.filter(sender=request.user).order_by('-created_at')
    return render(request, 'core/profile.html', {
        'user_files': user_files,
        'user_messages': user_messages,
    })


@login_required
def send_message_to_admin(request):
    """Пайдаланушыдан админге хабарлама жіберу"""
    from django.contrib import messages
    from .models import UserMessage
    if request.method == 'POST':
        subject = request.POST.get('subject', '').strip()
        message_text = request.POST.get('message', '').strip()
        if subject and message_text:
            UserMessage.objects.create(
                sender=request.user,
                subject=subject,
                message=message_text,
            )
            messages.success(request, "Хабарламаңыз жіберілді!")
        else:
            messages.error(request, "Тақырып пен хабарлама мәтінін толтырыңыз.")
    return redirect('profile')


@login_required
def update_disability(request):
    from .models import UserProfile
    if request.method == 'POST':
        disability_type = request.POST.get('disability_type', 'visual')
        if disability_type not in ('visual', 'hearing'):
            disability_type = 'visual'
        profile, _ = UserProfile.objects.get_or_create(user=request.user)
        profile.disability_type = disability_type
        profile.save()
        messages.success(request, "Профиль сәтті жаңартылды!")
    return redirect('profile')


def about(request):
    return render(request, 'core/about.html')


@login_required
def settings_view(request):
    return render(request, 'core/settings.html')

@login_required
def speech_to_text(request):
    return render(request, 'core/speech_to_text.html')


@login_required
def transcribe_mic(request):
    """Браузерден жазылған аудио → Whisper арқылы мәтін"""
    import json
    if request.method != 'POST':
        return HttpResponse(status=405)

    audio_blob = request.FILES.get('audio')
    lang_code = request.POST.get('lang', 'kk')  # kk / ru / en

    if not audio_blob:
        return HttpResponse(
            json.dumps({'error': 'Аудио жіберілмеді'}, ensure_ascii=False),
            content_type='application/json', status=400
        )

    os.makedirs('media/uploads', exist_ok=True)
    # webm/ogg blob — Whisper ffmpeg арқылы оқи алады
    ext = '.webm'
    audio_path = os.path.join('media/uploads', f'mic_tmp_{request.user.id}{ext}')
    with open(audio_path, 'wb+') as f:
        for chunk in audio_blob.chunks():
            f.write(chunk)

    whisper_lang_map = {'kk-KZ': 'kk', 'ru-RU': 'ru', 'en-US': 'en'}
    whisper_lang = whisper_lang_map.get(lang_code, lang_code)

    try:
        segments_gen, _ = _transcribe_fast(audio_path, language=whisper_lang)
        text = ' '.join(s.text.strip() for s in segments_gen).strip()
        return HttpResponse(
            json.dumps({'text': text}, ensure_ascii=False),
            content_type='application/json'
        )
    except Exception as e:
        return HttpResponse(
            json.dumps({'error': str(e)}, ensure_ascii=False),
            content_type='application/json', status=500
        )
    finally:
        try:
            os.remove(audio_path)
        except Exception:
            pass


# ─────────────────────────────────────────────
# 6. AUDIO TEXT — Аудио → Мәтін (Whisper)
# ─────────────────────────────────────────────
def convert_audio_to_text(myfile, user):
    """MP3/WAV/M4A/OGG аудио → мәтін (Whisper)"""
    os.makedirs('media/uploads', exist_ok=True)
    os.makedirs('media/ocr', exist_ok=True)

    audio_path = os.path.join('media/uploads', myfile.name)
    with open(audio_path, 'wb+') as f:
        for chunk in myfile.chunks():
            f.write(chunk)

    lang_names = {
        'kk': 'Қазақша', 'ru': 'Орысша', 'en': 'Ағылшынша', 'unknown': 'Белгісіз'
    }
    try:
        segments_gen, info = _transcribe_fast(audio_path)
        text = ' '.join(s.text.strip() for s in segments_gen).strip()
        detected_lang = info.language or 'unknown'

        if not text:
            return None, None, None, "Аудиодан мәтін табылмады"

        txt_name = f"audio_{os.path.splitext(myfile.name)[0]}.txt"
        txt_path = os.path.join('media/ocr', txt_name)
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)

        ConvertedFile.objects.create(
            user=user,
            title=myfile.name,
            conversion_type='audio_text',
            converted_text=text[:1000],
        )
        return text, f"/media/ocr/{txt_name}", lang_names.get(detected_lang, detected_lang.upper()), None
    except Exception as e:
        return None, None, None, str(e)
    finally:
        try:
            os.remove(audio_path)
        except Exception:
            pass


def admin_required(view_func):
    """Тек админдерге рұқсат"""
    from functools import wraps
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated or not request.user.is_superuser:
            return redirect('home')
        return view_func(request, *args, **kwargs)
    return wrapper


@admin_required
def admin_dashboard(request):
    """Админ панелінің басты беті"""
    from django.contrib.auth.models import User
    from datetime import datetime, timedelta
    from django.utils import timezone

    today = timezone.now().date()
    week_ago = timezone.now() - timedelta(days=7)

    # Статистика
    stats = {
        'total_users': User.objects.count(),
        'total_files': ConvertedFile.objects.count(),
        'new_users_today': User.objects.filter(date_joined__date=today).count(),
        'files_today': ConvertedFile.objects.filter(created_at__date=today).count(),
        'tts_count': ConvertedFile.objects.filter(conversion_type='tts').count(),
        'subtitle_count': ConvertedFile.objects.filter(conversion_type='subtitles').count(),
        'ocr_count': ConvertedFile.objects.filter(conversion_type='ocr').count(),
        'large_text_count': ConvertedFile.objects.filter(conversion_type='large_text').count(),
    }

    # Апталық деректер
    weekly_data = []
    for i in range(6, -1, -1):
        day = timezone.now() - timedelta(days=i)
        count = ConvertedFile.objects.filter(created_at__date=day.date()).count()
        weekly_data.append(count)

    # Жүйе ақпараты
    import sys, django as dj
    system_info = {
        'django_version': dj.get_version(),
        'python_version': sys.version.split()[0],
        'media_size': get_media_size(),
    }

    # Жүйе параметрлері
    from .models import SystemSettings, Notification, UserMessage
    system_settings, _ = SystemSettings.objects.get_or_create(pk=1)
    notifications = Notification.objects.all()[:10]
    user_messages = UserMessage.objects.all().order_by('-created_at')
    new_messages_count = UserMessage.objects.filter(status='new').count()

    return render(request, 'core/admin_dashboard.html', {
        'users': User.objects.all().order_by('-date_joined'),
        'all_files': ConvertedFile.objects.all().order_by('-created_at')[:50],
        'stats': stats,
        'weekly_data': weekly_data,
        'system_info': system_info,
        'system_settings': system_settings,
        'notifications': notifications,
        'user_messages': user_messages,
        'new_messages_count': new_messages_count,
    })


def get_media_size():
    """Media папкасының өлшемін есептеу"""
    total = 0
    try:
        for dirpath, dirnames, filenames in os.walk('media'):
            for f in filenames:
                fp = os.path.join(dirpath, f)
                total += os.path.getsize(fp)
        mb = total / (1024 * 1024)
        return f"{mb:.1f} МБ"
    except:
        return "0 МБ"


@admin_required
def admin_toggle_user(request, user_id):
    """Пайдаланушыны белсендіру/блоктау"""
    from django.contrib.auth.models import User
    from django.contrib import messages
    if request.method == 'POST':
        user = get_object_or_404(User, id=user_id)
        user.is_active = not user.is_active
        user.save()
        status = "белсендірілді" if user.is_active else "блокталды"
        messages.success(request, f"{user.username} {status}")
    return redirect('admin_dashboard')


@admin_required
def admin_delete_user(request, user_id):
    """Пайдаланушыны жою"""
    from django.contrib.auth.models import User
    from django.contrib import messages
    if request.method == 'POST':
        user = get_object_or_404(User, id=user_id)
        username = user.username
        user.delete()
        messages.success(request, f"{username} жойылды")
    return redirect('admin_dashboard')


@admin_required
def admin_delete_file(request, file_id):
    """Файлды жою"""
    from django.contrib import messages
    if request.method == 'POST':
        file = get_object_or_404(ConvertedFile, id=file_id)
        title = file.title
        file.delete()
        messages.success(request, f"{title} жойылды")
    return redirect('admin_dashboard')


@admin_required
def admin_send_notification(request):
    """Хабарлама жіберу"""
    from django.contrib.auth.models import User
    from django.contrib import messages
    from .models import Notification
    if request.method == 'POST':
        subject = request.POST.get('subject', '')
        message = request.POST.get('message', '')
        recipient_id = request.POST.get('recipient', 'all')

        if recipient_id == 'all':
            Notification.objects.create(
                subject=subject,
                message=message,
                recipient=None
            )
            messages.success(request, "Хабарлама барлық пайдаланушыларға жіберілді!")
        else:
            user = get_object_or_404(User, id=recipient_id)
            Notification.objects.create(
                subject=subject,
                message=message,
                recipient=user
            )
            messages.success(request, f"Хабарлама {user.username}-ға жіберілді!")

    return redirect('admin_dashboard')


@admin_required
def admin_mark_message(request, message_id):
    """Хабарлама статусын өзгерту"""
    from .models import UserMessage
    if request.method == 'POST':
        msg = get_object_or_404(UserMessage, id=message_id)
        msg.status = request.POST.get('status', 'read')
        msg.save()
    return redirect('admin_dashboard')


@admin_required
def admin_reply_message(request, message_id):
    """Пайдаланушыға жауап жіберу"""
    from .models import UserMessage, Notification
    from django.utils import timezone
    from django.contrib import messages
    if request.method == 'POST':
        msg = get_object_or_404(UserMessage, id=message_id)
        reply_text = request.POST.get('reply', '').strip()
        if reply_text:
            msg.reply = reply_text
            msg.replied_at = timezone.now()
            msg.status = 'replied'
            msg.save()
            messages.success(request, f"{msg.sender.username}-ға жауап жіберілді!")
    return redirect('admin_dashboard')


@admin_required
def admin_save_settings(request):
    """Жүйе параметрлерін сақтау"""
    from django.contrib import messages
    from .models import SystemSettings
    if request.method == 'POST':
        settings_obj, _ = SystemSettings.objects.get_or_create(pk=1)
        settings_obj.tts_voice = request.POST.get('tts_voice', 'kk-KZ-AigulNeural')
        settings_obj.default_lang = request.POST.get('default_lang', 'kk')
        settings_obj.whisper_model = request.POST.get('whisper_model', 'base')
        settings_obj.max_file_size = int(request.POST.get('max_file_size', 100))
        settings_obj.save()
        messages.success(request, "Параметрлер сақталды!")
    return redirect('admin_dashboard')
